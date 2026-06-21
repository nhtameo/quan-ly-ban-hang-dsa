from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT_DIR = Path(__file__).resolve().parents[1]
REPORT_MD = ROOT_DIR / "report" / "report.md"
REPORT_DOCX = ROOT_DIR / "report" / "report.docx"


def set_run_font(run, name: str = "Calibri", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "CodeBlock" not in styles:
        code_style = styles.add_style("CodeBlock", 1)
    else:
        code_style = styles["CodeBlock"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, text in enumerate(rows[0]):
        header_cells[index].text = text
        header_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cells[index].text = text
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [part.strip() for part in stripped.split("|")]


def build_docx() -> Path:
    doc = Document()
    configure_document(doc)
    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()

    i = 0
    in_code = False
    code_lines: list[str] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            text = " ".join(part.strip() for part in paragraph_lines).strip()
            if text:
                doc.add_paragraph(text)
            paragraph_lines.clear()

    def flush_code() -> None:
        if code_lines:
            for code_line in code_lines:
                p = doc.add_paragraph(style="CodeBlock")
                p.add_run(code_line)
            code_lines.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("| ---"):
            flush_paragraph()
            table_rows = [parse_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_markdown_table(doc, table_rows)
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:].strip())
            set_run_font(run, size=22, bold=True)
            i += 1
            continue
        if stripped.startswith("## "):
            flush_paragraph()
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            flush_paragraph()
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("- "):
            flush_paragraph()
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            i += 1
            continue
        if stripped == "---":
            flush_paragraph()
            i += 1
            continue

        paragraph_lines.append(line)
        i += 1

    flush_paragraph()
    REPORT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_DOCX)
    return REPORT_DOCX


if __name__ == "__main__":
    out = build_docx()
    print(out)

